from typing import List, Dict, Any
from .base import BaseChunkStrategy
from docx import Document
import logging
import os
import traceback

# 设置日志
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class WordChunkStrategy(BaseChunkStrategy):
    """Word文档切块策略实现 - 只实现 chunk_with_meta 方法"""
    
    def __init__(self):
        super().__init__()  # 必须调用父类初始化
        self.current_heading = ""
    
    def chunk_with_meta(self, file_path: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """
        按照指定大小切分Word文档，返回包含内容和元数据的结果
        
        Args:
            file_path: Word文档路径
            chunk_size: 每个块的大小（字符数）
            overlap: 相邻块之间的重叠字符数
            
        Returns:
            包含内容和元数据的文本块列表
        """
        try:
            logger.info(f"开始处理Word文档: {file_path}, chunk_size={chunk_size}, overlap={overlap}")
            
            # 检查文件类型
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext != '.docx':
                raise ValueError(f"不支持的文件类型: {file_ext}。Word策略只支持.docx文件，请为{file_ext}文件选择合适的策略，如'.txt'应选择'纯文本切块策略'")
            
            # 加载Word文档
            logger.info("正在加载Word文档...")
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            logger.info(f"Word文档加载成功，总段落数: {total_paragraphs}")
            
            # 初始化结果
            result_chunks = []
            current_chunk = []
            current_size = 0
            self.current_heading = ""
            
            # 遍历段落
            for i, para in enumerate(doc.paragraphs):
                try:
                    # 跳过空段落
                    if not para.text.strip():
                        continue
                    
                    # 检查是否是标题
                    is_heading = False
                    try:
                        if para.style.name.startswith('Heading'):
                            is_heading = True
                            self.current_heading = para.text
                            logger.debug(f"发现标题: {self.current_heading}")
                            
                            # 如果当前chunk不为空，先输出
                            if current_chunk:
                                result_chunks.append(self._create_chunk(current_chunk, self.current_heading))
                                current_chunk = []
                                current_size = 0
                            
                            # 标题单独成一个块
                            continue
                    except Exception as e:
                        logger.debug(f"样式检测失败: {str(e)}")
                        pass  # 样式检测失败，当作普通段落处理
                    
                    # 处理普通段落
                    text = para.text
                    text_length = len(text)
                    
                    # 如果添加这个段落会超过chunk_size，且当前chunk不为空，先输出当前chunk
                    if current_size + text_length > chunk_size and current_chunk:
                        result_chunks.append(self._create_chunk(current_chunk, self.current_heading))
                        current_chunk = []
                        current_size = 0
                    
                    # 如果段落本身就超过chunk_size，分段处理
                    if text_length > chunk_size:
                        # 处理长段落
                        try:
                            para_chunks = self._chunk_long_text(text, chunk_size, overlap)
                            for p_chunk in para_chunks:
                                result_chunks.append(self._create_chunk([p_chunk], self.current_heading))
                        except Exception as e:
                            logger.error(f"处理段落 {i} 时出错: {str(e)}")
                            # 如果处理失败，将整个段落作为一个块
                            result_chunks.append(self._create_chunk([text], self.current_heading))
                    else:
                        # 正常添加段落
                        current_chunk.append(text)
                        current_size += text_length
                except Exception as e:
                    logger.error(f"处理段落 {i} 时发生错误: {str(e)}")
                    logger.error(traceback.format_exc())
                    continue  # 继续处理下一个段落
            
            # 处理最后一个chunk
            if current_chunk:
                result_chunks.append(self._create_chunk(current_chunk, self.current_heading))
            
            logger.info(f"文档处理完成，共生成 {len(result_chunks)} 个文本块")
            return result_chunks
            
        except Exception as e:
            logger.error(f"切块过程中发生错误: {str(e)}")
            logger.error(traceback.format_exc())
            raise
    
    def _chunk_long_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        将长文本分割成多个块
        
        Args:
            text: 长文本
            chunk_size: 每个块的大小
            overlap: 重叠大小
            
        Returns:
            分割后的文本块列表
        """
        chunks = []
        start = 0
        text_length = len(text)
        last_end = -1  # 用于检测是否卡住
        
        while start < text_length:
            # 安全检查：如果start没有前进，说明可能卡住了
            if start == last_end:
                logger.warning(f"检测到可能的死循环，强制结束处理。当前位置: {start}, 文本长度: {text_length}")
                # 强制添加剩余文本
                if start < text_length:
                    chunks.append(text[start:])
                break
                
            last_end = start
            
            # 计算当前块的结束位置
            end = min(start + chunk_size, text_length)
            
            # 避免在单词中间切断
            if end < text_length and not text[end].isspace():
                # 向后查找空格
                space_pos = text.find(' ', end)
                if space_pos != -1 and space_pos - end < 20:  # 最多向后查找20个字符
                    end = space_pos
            
            # 添加当前块
            chunks.append(text[start:end])
            
            # 考虑重叠
            start = end - overlap if overlap > 0 else end
            
            # 确保start至少前进一个字符
            if start <= last_end:
                start = last_end + 1
        
        logger.debug(f"长文本分割完成，共生成 {len(chunks)} 个子块")
        return chunks
    
    def _create_chunk(self, texts: List[str], heading: str) -> Dict[str, Any]:
        """
        创建包含内容和元数据的块
        
        Args:
            texts: 文本列表
            heading: 当前标题
            
        Returns:
            包含内容和元数据的字典
        """
        return {
            "content": "\n".join(texts),
            "meta": {
                "heading": heading,
                "char_count": sum(len(t) for t in texts)
            }
        }
    
    def get_metadata(self) -> Dict[str, Any]:
        """
        获取切块策略的元数据
        
        Returns:
            包含策略信息的字典
        """
        return {
            "name": "word",
            "display_name": "Word文档切块策略",
            "description": "Word文档切块策略，基于标题和指定大小进行切块",
            "supported_types": [".docx"]
        } 