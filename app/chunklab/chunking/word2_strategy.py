from typing import List, Dict, Any, Tuple
from .base import BaseChunkStrategy
from docx import Document
import re
import json
import logging
import os
import traceback

# 设置日志
logging.basicConfig(level=logging.INFO)  # 将日志级别改为INFO，减少输出
logger = logging.getLogger(__name__)

class Word2ChunkStrategy(BaseChunkStrategy):
    """Word文档智能标题合并切块策略 - 基于标题树结构动态生成chunks"""
    
    def __init__(self):
        super().__init__()  # 必须调用父类初始化
        self.title_regex = re.compile(r'^(\d+(?:\.\d+)*)\s+(.*)$')
    
    def chunk_with_meta(self, file_path: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """
        按照标题树结构智能切分Word文档，返回包含内容和元数据的结果
        
        Args:
            file_path: Word文档路径
            chunk_size: 每个块的大小（字符数）- 用于处理过长内容
            overlap: 相邻块之间的重叠字符数
            
        Returns:
            包含内容和元数据的文本块列表
        """
        try:
            logger.info(f"开始处理Word文档: {file_path}, chunk_size={chunk_size}, overlap={overlap}")
            
            # 检查文件类型
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext != '.docx':
                raise ValueError(f"不支持的文件类型: {file_ext}。Word2策略只支持.docx文件")
            
            # 加载Word文档
            logger.info("正在加载Word文档...")
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            logger.info(f"Word文档加载成功，总段落数: {total_paragraphs}")
            
            # 获取文档中所有非空段落
            lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            logger.info(f"提取到非空段落数: {len(lines)}")
            
            # 构建标题树和内容块
            chunks = self._parse_docx_to_chunks(lines, chunk_size, overlap)
            
            logger.info(f"文档处理完成，共生成 {len(chunks)} 个文本块")
            return chunks
            
        except Exception as e:
            logger.error(f"切块过程中发生错误: {str(e)}")
            logger.error(traceback.format_exc())
            raise
    
    def _parse_docx_to_chunks(self, lines: List[str], chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """
        解析文档内容，构建标题树并生成智能块
        
        Args:
            lines: 文档的非空行列表
            chunk_size: 最大块大小
            overlap: 重叠大小
            
        Returns:
            切分后的块列表
        """
        # 直接返回简化版处理结果，避免卡死问题
        chunks = []
        stack = []  # 当前标题栈
        buffer = []  # 当前内容缓冲区
        
        for line in lines:
            match = self.title_regex.match(line)
            
            if match:
                # 发现标题
                title_num = match.group(1)
                title_text = match.group(2)
                current_level = title_num.count('.') + 1  # 计算标题层级
                
                # 如果buffer不为空，先生成chunk
                if buffer:
                    chunk_text = "\n".join(buffer)
                    chunks.append({
                        "content": chunk_text,
                        "meta": {
                            "title_path": json.dumps(list(stack)) if stack else json.dumps(["文档开头"])
                        }
                    })
                    buffer = []
                
                # 调整标题栈到当前层级
                while len(stack) > 0 and len(stack) >= current_level:
                    stack.pop()
                stack.append(f"{title_num} {title_text}")
            else:
                # 普通内容
                buffer.append(line)
                
                # 如果buffer太大，直接生成chunk
                if len("\n".join(buffer)) > chunk_size:
                    chunk_text = "\n".join(buffer)
                    chunks.append({
                        "content": chunk_text,
                        "meta": {
                            "title_path": json.dumps(list(stack)) if stack else json.dumps(["文档开头"])
                        }
                    })
                    buffer = []
        
        # 处理最后的缓冲区
        if buffer:
            chunk_text = "\n".join(buffer)
            chunks.append({
                "content": chunk_text,
                "meta": {
                    "title_path": json.dumps(list(stack)) if stack else json.dumps(["文档开头"])
                }
            })
        
        return chunks
    
    def _chunk_long_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        将长文本分割成多个块
        
        Args:
            text: 长文本
            chunk_size: 每个块的大小
            overlap: 重叠大小
            
        Returns:
            分割后的文本块列表
        """
        chunks = []
        start = 0
        text_length = len(text)
        last_position = -1  # 用于检测循环是否前进
        
        while start < text_length:
            # 安全检查：避免无限循环
            if start == last_position:
                logger.warning(f"检测到可能的死循环，在位置 {start}，强制前进")
                start += 1
                continue
                
            last_position = start
            
            # 计算当前块的结束位置
            end = min(start + chunk_size, text_length)
            
            # 避免在单词中间切断
            if end < text_length and not text[end].isspace():
                # 向后查找空格
                space_pos = text.find(' ', end)
                if space_pos != -1 and space_pos - end < 20:  # 最多向后查找20个字符
                    end = space_pos + 1  # 包含空格
            
            # 添加当前块
            chunks.append(text[start:end])
            
            # 考虑重叠
            start = end - overlap if overlap > 0 and end - overlap > start else end
        
        logger.info(f"长文本分割完成，共生成 {len(chunks)} 个子块")
        return chunks
    
    def get_metadata(self) -> Dict[str, Any]:
        """
        获取切块策略的元数据
        
        Returns:
            包含策略信息的字典
        """
        return {
            "name": "word2",
            "display_name": "Word智能标题合并切块策略",
            "description": "基于标题树结构，智能合并无内容标题，动态生成文本块",
            "supported_types": [".docx"]
        } 